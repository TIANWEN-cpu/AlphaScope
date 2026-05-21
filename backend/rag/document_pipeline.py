"""
Document Pipeline: 文档处理管线。

架构文档要求：Upload → Virus Scan → Parse → Clean → Chunk → Embedding → Index → Retrieval → Citation

支持格式：PDF、DOCX、XLSX、CSV、TXT、Markdown
"""

import hashlib
import logging
import time
from pathlib import Path
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = {".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md", ".json"}


@dataclass
class ProcessedDocument:
    """处理后的文档"""

    doc_id: str
    filename: str
    file_type: str
    content: str = ""
    chunks: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunk_count: int = 0
    processing_time_ms: float = 0


class DocumentPipeline:
    """文档处理管线"""

    def __init__(self):
        self._processed: Dict[str, ProcessedDocument] = {}

    def process_file(
        self, file_path: str, metadata: Optional[Dict] = None
    ) -> Optional[ProcessedDocument]:
        """处理文件：解析 → 清洗 → 分块 → 索引"""
        t0 = time.time()
        p = Path(file_path)

        if not p.exists():
            logger.warning(f"文件不存在: {file_path}")
            return None

        suffix = p.suffix.lower()
        if suffix not in SUPPORTED_FORMATS:
            logger.warning(f"不支持的文件格式: {suffix}")
            return None

        # 解析
        content = self._parse_file(p, suffix)
        if not content:
            return None

        # 清洗
        content = self._clean_content(content)

        # 分块
        chunks = self._chunk_content(content)

        # 生成文档 ID
        doc_id = hashlib.md5(f"{p.name}:{len(content)}".encode()).hexdigest()[:12]

        doc = ProcessedDocument(
            doc_id=doc_id,
            filename=p.name,
            file_type=suffix,
            content=content,
            chunks=chunks,
            metadata=metadata or {},
            chunk_count=len(chunks),
            processing_time_ms=round((time.time() - t0) * 1000, 1),
        )

        self._processed[doc_id] = doc

        # 索引到 RAG
        self._index_document(doc)

        return doc

    def process_text(
        self, text: str, filename: str = "text", metadata: Optional[Dict] = None
    ) -> ProcessedDocument:
        """处理文本内容"""
        t0 = time.time()

        content = self._clean_content(text)
        chunks = self._chunk_content(content)
        doc_id = hashlib.md5(f"{filename}:{len(content)}".encode()).hexdigest()[:12]

        doc = ProcessedDocument(
            doc_id=doc_id,
            filename=filename,
            file_type="text",
            content=content,
            chunks=chunks,
            metadata=metadata or {},
            chunk_count=len(chunks),
            processing_time_ms=round((time.time() - t0) * 1000, 1),
        )

        self._processed[doc_id] = doc
        self._index_document(doc)
        return doc

    def _parse_file(self, path: Path, suffix: str) -> str:
        """解析文件内容"""
        try:
            if suffix in (".txt", ".md", ".json", ".csv"):
                return path.read_text(encoding="utf-8", errors="ignore")

            elif suffix == ".pdf":
                return self._parse_pdf(path)

            elif suffix == ".docx":
                return self._parse_docx(path)

            elif suffix == ".xlsx":
                return self._parse_xlsx(path)

        except Exception as e:
            logger.warning(f"解析文件失败 {path.name}: {e}")

        return ""

    def _parse_pdf(self, path: Path) -> str:
        """解析 PDF"""
        try:
            import subprocess

            result = subprocess.run(
                ["pdftotext", str(path), "-"],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0:
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Fallback: 尝试 PyPDF2
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            text_parts = []
            for page in reader.pages[:50]:  # 最多 50 页
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except ImportError:
            pass

        return ""

    def _parse_docx(self, path: Path) -> str:
        """解析 DOCX"""
        try:
            import docx

            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            return ""
        except Exception:
            return ""

    def _parse_xlsx(self, path: Path) -> str:
        """解析 XLSX"""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(str(path), read_only=True)
            text_parts = []
            for sheet in wb.worksheets[:5]:  # 最多 5 个 sheet
                for row in sheet.iter_rows(
                    max_row=100, values_only=True
                ):  # 每 sheet 最多 100 行
                    row_text = " | ".join(str(c) for c in row if c is not None)
                    if row_text.strip():
                        text_parts.append(row_text)
            return "\n".join(text_parts)
        except ImportError:
            return ""
        except Exception:
            return ""

    def _clean_content(self, content: str) -> str:
        """清洗内容"""
        # 去除多余空白
        lines = content.splitlines()
        cleaned = []
        for line in lines:
            line = line.strip()
            if line:
                cleaned.append(line)
        return "\n".join(cleaned)

    def _chunk_content(
        self, content: str, chunk_size: int = 500, overlap: int = 50
    ) -> List[str]:
        """分块内容"""
        if len(content) <= chunk_size:
            return [content]

        chunks = []
        start = 0
        while start < len(content):
            end = start + chunk_size
            # 尝试在句子边界分割
            if end < len(content):
                # 找最近的句号/换行
                for sep in ["\n", "。", ".", "！", "!", "？", "?"]:
                    pos = content.rfind(sep, start + chunk_size // 2, end)
                    if pos > start:
                        end = pos + 1
                        break

            chunk = content[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap
            if start >= len(content):
                break

        return chunks

    def _index_document(self, doc: ProcessedDocument):
        """将文档索引到 RAG"""
        try:
            from backend.rag.vector_store import VectorStore

            store = VectorStore()
            ids = [f"{doc.doc_id}_chunk_{i}" for i in range(len(doc.chunks))]
            metadatas = [
                {
                    "doc_id": doc.doc_id,
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                    "chunk_index": i,
                    **doc.metadata,
                }
                for i in range(len(doc.chunks))
            ]
            store.add_documents(
                collection_name="user_documents",
                documents=doc.chunks,
                metadatas=metadatas,
                ids=ids,
            )
            logger.info(f"已索引文档 {doc.filename}: {doc.chunk_count} 个分块")
            return ids
        except Exception as e:
            logger.debug(f"索引文档失败: {e}")
            return []

    def process_and_persist(
        self, file_path: str, metadata: Optional[Dict] = None
    ) -> Optional[ProcessedDocument]:
        """处理文件 → 保存到 SQLite → 索引到向量库"""
        from backend.file_store import content_hash, save_chunks, save_document

        p = Path(file_path)
        if not p.exists():
            logger.warning(f"文件不存在: {file_path}")
            return None

        # 计算内容 hash
        try:
            raw = p.read_bytes()
            c_hash = content_hash(raw)
        except Exception:
            c_hash = ""

        # 处理文件（解析 → 清洗 → 分块）
        doc = self.process_file(file_path, metadata)
        if not doc:
            return None

        # 保存到 SQLite
        saved = save_document(
            title=doc.filename,
            file_path=file_path,
            content_hash=c_hash,
            source_type="upload",
            metadata=metadata or {},
        )
        doc.doc_id = saved["id"]

        # 保存 chunks 到 SQLite
        save_chunks(doc.doc_id, doc.chunks)

        # 索引到向量库（ChromaDB 不可用时跳过）
        self._index_document(doc)

        return doc

    def get_document(self, doc_id: str) -> Optional[ProcessedDocument]:
        """获取已处理的文档"""
        return self._processed.get(doc_id)

    def list_documents(self) -> List[Dict[str, Any]]:
        """列出已处理的文档"""
        return [
            {
                "doc_id": d.doc_id,
                "filename": d.filename,
                "type": d.file_type,
                "chunks": d.chunk_count,
                "processing_time_ms": d.processing_time_ms,
            }
            for d in self._processed.values()
        ]


# 单例
_pipeline: Optional[DocumentPipeline] = None


def get_document_pipeline() -> DocumentPipeline:
    """获取全局文档处理管线"""
    global _pipeline
    if _pipeline is None:
        _pipeline = DocumentPipeline()
    return _pipeline
